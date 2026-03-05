#!/usr/bin/env python3
"""Stride Deliverables — Word (.docx) Builder
=============================================
Generates native Word documents with proper <w:tbl> table objects:
  1. Stride Policy Cross Reference V7.docx
  2. Stride SRA V8.docx
  3. Stride ASR Questionnaire V4.docx
  4. Governance Memos (8 memos from markdown sources)

Requires: python-docx, openpyxl (for dates) in venv
"""

import math, os, re, json, glob, yaml
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent                    # <repo>/build/
REPO = BASE.parent                                        # <repo>/
EXTRACTED = BASE / "extracted"
PNG_DIR = REPO / "diagrams" / "rendering_png_v2"
OUTPUT = BASE / "output"
OUTPUT.mkdir(exist_ok=True)
MEMO_DIR = REPO / "workpapers"
TEMPLATE_DOCX = BASE / "stride_template.docx"

# ── Figure tracking ─────────────────────────────────────────────────────
_figure_counter = 0

def reset_figure_counter():
    """Reset figure counter for a new document."""
    global _figure_counter
    _figure_counter = 0

def new_doc():
    """Create a new Document from the template (has GridTable4-Accent1 style)."""
    reset_figure_counter()
    if TEMPLATE_DOCX.exists():
        return Document(str(TEMPLATE_DOCX))
    return Document()

# ── Load extracted data ───────────────────────────────────────────────
with open(EXTRACTED / "policies.json") as f:
    policies_data = json.load(f)
with open(EXTRACTED / "workpapers.json") as f:
    workpapers_data = json.load(f)

# ── Brand Colors ──────────────────────────────────────────────────────
STRIDE_GREEN = RGBColor(0x2E, 0x7D, 0x32)
STRIDE_BLUE  = RGBColor(0x15, 0x65, 0xC0)
STRIDE_GRAY  = RGBColor(0x75, 0x75, 0x75)
GAP_RED      = RGBColor(0xC6, 0x28, 0x28)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)

HEAD_GRAY_HEX = "E0E0E0"
CRIT_BG_HEX   = "FFEBEE"
HIGH_BG_HEX   = "FFF3E0"
MED_BG_HEX    = "E8F5E9"
CSF_BOX_HEX   = "E3F2FD"
FERPA_BOX_HEX = "FFF3E0"
SOX_BOX_HEX   = "F3E5F5"

# ── Helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_row_shading(row, hex_color):
    """Set background shading on all cells in a row."""
    for cell in row.cells:
        set_cell_shading(cell, hex_color)

def add_formatted_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run

def make_landscape(doc):
    """Set document to landscape orientation."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

def add_header_footer(doc, header_text, footer_date="March 3, 2026"):
    """Add header and footer to all sections."""
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        run = hp.add_run(header_text)
        run.font.size = Pt(8)
        run.font.color.rgb = STRIDE_GRAY
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        tab_run = hp.add_run("\tCONFIDENTIAL")
        tab_run.font.size = Pt(8)
        tab_run.font.color.rgb = STRIDE_GRAY

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        run = fp.add_run(footer_date)
        run.font.size = Pt(8)
        run.font.color.rgb = STRIDE_GRAY
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_title_page(doc, title, subtitle, extra_lines=None):
    """Add a branded title page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, title, bold=True, color=STRIDE_GREEN, size=Pt(28))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, subtitle, size=Pt(18), color=STRIDE_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Stride, Inc.", size=Pt(14))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Information Security & Governance", size=Pt(12), color=STRIDE_GRAY)

    if extra_lines:
        doc.add_paragraph("")
        for line in extra_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, line, size=Pt(10), color=STRIDE_GRAY)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, f"Document Date: March 3, 2026", size=Pt(10))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Classification: Confidential", size=Pt(10))

    doc.add_page_break()


def clean_text(text):
    """Remove hard line breaks from text so it wraps naturally in table cells."""
    if not text:
        return text
    # Replace newlines with spaces, collapse multiple spaces
    return re.sub(r'\s+', ' ', text.replace('\n', ' ')).strip()


def set_cell_text(cell, text):
    """Set cell text, stripping hard line breaks so content wraps with the cell."""
    cell.text = clean_text(text)


def ensure_copyable(doc):
    """Remove any document protection so text can be freely selected and copied."""
    try:
        for rel in doc.part.rels.values():
            tp = getattr(rel, 'target_part', None)
            if tp and hasattr(tp, 'element') and tp.element is not None:
                pn = str(getattr(tp, 'partname', ''))
                if 'settings' in pn.lower():
                    for dp in list(tp.element.iter(qn('w:documentProtection'))):
                        dp.getparent().remove(dp)
                    # Also set the default view to Print Layout (not Read Mode)
                    view_elem = tp.element.find(qn('w:view'))
                    if view_elem is not None:
                        view_elem.set(qn('w:val'), 'print')
                    else:
                        view_elem = parse_xml(f'<w:view {nsdecls("w")} w:val="print"/>')
                        tp.element.insert(0, view_elem)
    except Exception:
        pass


def add_figure(doc, image_path, title, width=Inches(8.5)):
    """Add a centered image with a numbered 'Figure X — Title' caption.

    Uses a Word SEQ field so the Table of Figures can collect captions.
    Returns the figure number assigned.
    """
    global _figure_counter
    _figure_counter += 1
    fig_num = _figure_counter

    # Centered image paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(image_path), width=width)

    # Caption paragraph — use built-in 'Caption' style if available, else Normal
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add SEQ field for auto-numbering (enables Table of Figures)
    # Format: "Figure " + SEQ field + " — Title"
    run_prefix = p_cap.add_run("Figure ")
    run_prefix.italic = True
    run_prefix.font.size = Pt(9)
    run_prefix.font.color.rgb = STRIDE_GRAY

    # Insert a SEQ Figure field
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" SEQ Figure \\\\* ARABIC ">'
        f'<w:r><w:rPr><w:i/><w:sz w:val="18"/><w:color w:val="757575"/></w:rPr>'
        f'<w:t>{fig_num}</w:t></w:r></w:fldSimple>'
    )
    fld_el = parse_xml(fld_xml)
    p_cap._element.append(fld_el)

    run_title = p_cap.add_run(f" \u2014 {title}")
    run_title.italic = True
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = STRIDE_GRAY

    # Set the paragraph style name for TOF collection (Word looks for 'Caption' style)
    try:
        p_cap.style = doc.styles['Caption']
    except KeyError:
        pass  # Will still work via SEQ field

    return fig_num


def add_table_of_contents(doc):
    """Insert a Table of Contents field (Word will populate it on first open/update)."""
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # TOC field: \o "1-3" = Heading 1-3, \h = hyperlinks, \z = hide tab leaders in web
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" TOC \\\\o &quot;1-3&quot; \\\\h \\\\z \\\\u ">'
        f'<w:r><w:rPr><w:i/><w:sz w:val="18"/></w:rPr>'
        f'<w:t>Update this field to see the Table of Contents (right-click \u2192 Update Field)</w:t>'
        f'</w:r></w:fldSimple>'
    )
    p._element.append(parse_xml(fld_xml))
    doc.add_page_break()


def add_table_of_figures(doc):
    """Insert a Table of Figures field (Word will populate it on first open/update)."""
    doc.add_heading("Table of Figures", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # TOC field scoped to SEQ Figure captions
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" TOC \\\\h \\\\z \\\\c &quot;Figure&quot; ">'
        f'<w:r><w:rPr><w:i/><w:sz w:val="18"/></w:rPr>'
        f'<w:t>Update this field to see the Table of Figures (right-click \u2192 Update Field)</w:t>'
        f'</w:r></w:fldSimple>'
    )
    p._element.append(parse_xml(fld_xml))
    doc.add_page_break()


def add_compliance_table(doc, sec, cross_refs=None):
    """Add a 3-column compliance alignment table for an SRA section.

    Columns: Framework | Control Target | Applicable References
    Always shows NIST CSF 2.0; shows FERPA and SOX rows only when applicable.
    Optional cross_refs: list of (label, value) tuples appended as additional rows.
    """
    if cross_refs is None:
        cross_refs = []
    # Count rows: header + NIST (always) + FERPA (if present) + SOX (if present) + cross-refs
    n_rows = 1 + 1  # header + NIST
    if "ferpa" in sec:
        n_rows += 1
    if "sox" in sec:
        n_rows += 1
    n_rows += len(cross_refs)

    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.rows[0].cells[0].text = "Framework"
    tbl.rows[0].cells[1].text = "Control Target"
    tbl.rows[0].cells[2].text = "Applicable References"

    row_idx = 1
    tbl.rows[row_idx].cells[0].text = "NIST CSF 2.0"
    tbl.rows[row_idx].cells[1].text = "Alignment"
    tbl.rows[row_idx].cells[2].text = ", ".join(sec["csf"])

    if "ferpa" in sec:
        row_idx += 1
        tbl.rows[row_idx].cells[0].text = "FERPA (34 CFR Part 99)"
        tbl.rows[row_idx].cells[1].text = "Relevance"
        tbl.rows[row_idx].cells[2].text = sec["ferpa"]

    if "sox" in sec:
        row_idx += 1
        tbl.rows[row_idx].cells[0].text = "SOX \u00a7404 ITGC"
        tbl.rows[row_idx].cells[1].text = "Relevance"
        tbl.rows[row_idx].cells[2].text = sec["sox"]

    for label, value in cross_refs:
        row_idx += 1
        tbl.rows[row_idx].cells[0].text = label
        tbl.rows[row_idx].cells[1].text = "Cross Reference"
        tbl.rows[row_idx].cells[2].text = value

    style_table(tbl)

    # Fixed column widths: Framework 1.5", Control Target 1.0", References = remainder
    for row in tbl.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(7.5)
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(1.0)
    tbl.columns[2].width = Inches(7.5)


def style_table(table, header_hex=None, numeric_cols=None):
    """Apply GridTable4-Accent1 style with full-width layout and auto row height.

    Uses the built-in Word alternating-row-shaded table style extracted from
    the baseline document.  Falls back to manual shading if style unavailable.
    numeric_cols: optional set of 0-based column indices that should be right-aligned.
    """
    if numeric_cols is None:
        numeric_cols = set()
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Apply GridTable4-Accent1 via XML (style name lookup doesn't work in python-docx)
    # Remove any existing tblStyle first
    existing_ts = tblPr.find(qn('w:tblStyle'))
    if existing_ts is not None:
        tblPr.remove(existing_ts)
    ts = parse_xml(f'<w:tblStyle {nsdecls("w")} w:val="GridTable4-Accent1"/>')
    tblPr.insert(0, ts)

    # Full-width table (100% of printable area)
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tblW)

    # Header row: bold + 9pt
    if table.rows:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
    # Body: 9pt font, auto row height
    for row_idx, row in enumerate(table.rows):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:hRule="auto" w:val="0"/>')
        trPr.append(trHeight)
        for col_idx, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            for p in cell.paragraphs:
                p.alignment = align
                if row_idx == 0:
                    continue
                for r in p.runs:
                    r.font.size = Pt(9)


# ── RSK/STORM Scoring (same as build_deliverables.py) ────────────────
RSK_DAMPING = 4
RSK_VMAX = 100
RSK_RAW_MAX = 134
WEIGHT_MAP = {"Critical": 4, "High": 3, "Medium": 2, "Info": 1}
WEIGHT_MAX = 4

def choice_scores(n):
    return [math.ceil(100 * i / n) for i in range(1, n + 1)]

def rsk_aggregate(measurements, a=RSK_DAMPING):
    valid = sorted((v for v in measurements if isinstance(v, (int, float)) and v > 0), reverse=True)
    if not valid:
        return 0
    return math.ceil(sum(v / (a ** j) for j, v in enumerate(valid)))

def rsk_normalize(raw, maximum=RSK_RAW_MAX):
    if maximum <= 0:
        return 0.0
    return min(100.0, raw / maximum * 100)


# ══════════════════════════════════════════════════════════════════════
# Import data structures from build_deliverables (inline copies to avoid
# import issues with the LaTeX-specific code)
# ══════════════════════════════════════════════════════════════════════

# ── NIST CSF 2.0 (abbreviated for cross-ref table) ───────────────────
NIST_CSF_FUNCTIONS = {
    "GV": "GOVERN",
    "ID": "IDENTIFY",
    "PR": "PROTECT",
    "DE": "DETECT",
    "RS": "RESPOND",
    "RC": "RECOVER",
}

# ── FERPA Key Sections ───────────────────────────────────────────────
FERPA_SECTIONS = {
    "§99.1": "Purpose — To set forth requirements for the protection of privacy of parents and students under FERPA",
    "§99.2": "Applicability — Applies to all educational agencies receiving DOE funding",
    "§99.3": "Definitions — Education records, directory information, personally identifiable information (PII), legitimate educational interest",
    "§99.4": "Student rights — Transfer of rights at age 18 or postsecondary enrollment",
    "§99.5": "Right of access — Right to inspect and review education records; 45-day compliance window",
    "§99.7": "Annual notification — Schools must annually notify parents/students of their FERPA rights",
    "§99.10": "Right to inspect — Parents/eligible students may inspect and review education records",
    "§99.12": "Right to amend — Right to request amendment of records believed to be inaccurate",
    "§99.20": "Amendment hearing — Hearing procedures when amendment is refused",
    "§99.30": "Prior consent — Written consent required before disclosing PII from education records",
    "§99.31": "Consent exceptions — Exceptions including school officials with legitimate educational interest, directory info, health/safety emergencies, judicial orders",
    "§99.32": "Record of disclosures — Institutions must maintain records of each request for and disclosure of PII",
    "§99.33": "Redisclosure limitations — Recipients may not redisclose PII without consent",
    "§99.34": "Transfer conditions — Conditions for disclosure to other schools",
    "§99.35": "Research exceptions — Disclosure for studies for, or on behalf of, educational agencies",
    "§99.36": "Health/safety emergency — Disclosure permitted in health or safety emergencies",
    "§99.37": "Directory information — Designation, opt-out rights, and de-identification requirements",
    "§99.60-67": "Enforcement — Investigation and enforcement procedures by the Family Policy Compliance Office",
}

# ── SOX 404 ITGC Domains ─────────────────────────────────────────────
SOX_404_ITGCS = {
    "Access Controls": {
        "description": "Restrict system access to authorized individuals; enforce least privilege and separation of duties",
        "controls": [
            "AC-1: Logical access provisioning/deprovisioning aligned with HR lifecycle events",
            "AC-2: Periodic user access reviews (quarterly for privileged, semi-annual for standard)",
            "AC-3: Privileged access management — just-in-time elevation, session recording, break-glass procedures",
            "AC-4: Authentication controls — MFA for all administrative and financially-significant system access",
            "AC-5: Service account governance — inventory, ownership, credential rotation",
        ]
    },
    "Change Management": {
        "description": "Ensure changes to IT systems are authorized, tested, and approved before production deployment",
        "controls": [
            "CM-1: Change request documentation including business justification and impact assessment",
            "CM-2: Segregation of development, test, and production environments",
            "CM-3: Code review / peer review before production deployment",
            "CM-4: Change approval by designated change authority (CAB or equivalent)",
            "CM-5: Emergency change procedures with post-implementation review",
            "CM-6: Release management and deployment automation controls",
        ]
    },
    "Computer Operations": {
        "description": "Ensure reliable processing of financial data through job scheduling, backup, and monitoring",
        "controls": [
            "CO-1: Job scheduling controls — automated scheduling with failure alerting",
            "CO-2: Backup and recovery — automated backups with tested restoration procedures",
            "CO-3: Incident management — defined escalation and resolution procedures",
            "CO-4: Batch processing integrity — reconciliation and completeness checks",
        ]
    },
    "Program Development": {
        "description": "Ensure systems are developed, tested, and implemented with appropriate controls",
        "controls": [
            "PD-1: System development lifecycle (SDLC) methodology with security requirements",
            "PD-2: Security testing — SAST, DAST, SCA integrated into CI/CD pipeline",
            "PD-3: User acceptance testing (UAT) for financially-significant system changes",
            "PD-4: Data migration controls with reconciliation and validation",
        ]
    },
}

# ── Policy Inventory ──────────────────────────────────────────────────
POLICY_INVENTORY = [
    {"id": "ISP 1.0", "title": "Information Security Program Policy", "csf_function": "GV", "csf_categories": ["GV.OC", "GV.RM", "GV.PO"], "ferpa": ["§99.1", "§99.2"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 1.1", "title": "IT Security Policy Index", "csf_function": "GV", "csf_categories": ["GV.PO"], "ferpa": [], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 1.2", "title": "IT Security Policy Definitions", "csf_function": "GV", "csf_categories": ["GV.PO", "GV.OC"], "ferpa": ["§99.3"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 2.1", "title": "Organizational Context (GV.OC)", "csf_function": "GV", "csf_categories": ["GV.OC"], "ferpa": ["§99.2"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 2.2", "title": "Risk Management Strategy (GV.RM)", "csf_function": "GV", "csf_categories": ["GV.RM"], "ferpa": [], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 2.3", "title": "Roles, Responsibilities, and Authorities (GV.RR)", "csf_function": "GV", "csf_categories": ["GV.RR"], "ferpa": [], "sox_itgc": ["Access Controls"], "type": "ISP"},
    {"id": "ISP 2.4", "title": "Policies, Processes, and Procedures (GV.PO)", "csf_function": "GV", "csf_categories": ["GV.PO"], "ferpa": ["§99.7"], "sox_itgc": ["Change Management"], "type": "ISP"},
    {"id": "ISP 2.5", "title": "Oversight (GV.OV)", "csf_function": "GV", "csf_categories": ["GV.OV"], "ferpa": ["§99.60-67"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 2.6", "title": "Cybersecurity Supply Chain Risk Management (GV.SC)", "csf_function": "GV", "csf_categories": ["GV.SC"], "ferpa": ["§99.33", "§99.35"], "sox_itgc": ["Access Controls"], "type": "ISP"},
    {"id": "ISP 3.1", "title": "Asset Management (ID.AM)", "csf_function": "ID", "csf_categories": ["ID.AM"], "ferpa": ["§99.3"], "sox_itgc": ["Computer Operations"], "type": "ISP"},
    {"id": "ISP 3.2", "title": "Risk Assessment (ID.RA)", "csf_function": "ID", "csf_categories": ["ID.RA"], "ferpa": ["§99.36"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 3.3", "title": "Response Improvement (ID.IM)", "csf_function": "ID", "csf_categories": ["ID.IM"], "ferpa": [], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 4.1", "title": "Identity Management, Authentication, and Access Control (PR.AA)", "csf_function": "PR", "csf_categories": ["PR.AA"], "ferpa": ["§99.31", "§99.32"], "sox_itgc": ["Access Controls"], "type": "ISP"},
    {"id": "ISP 4.2", "title": "Awareness and Training Standards (PR.AT)", "csf_function": "PR", "csf_categories": ["PR.AT"], "ferpa": ["§99.7"], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 4.3", "title": "Data Security (PR.DS)", "csf_function": "PR", "csf_categories": ["PR.DS"], "ferpa": ["§99.30", "§99.31", "§99.33"], "sox_itgc": ["Access Controls", "Computer Operations"], "type": "ISP"},
    {"id": "ISP 4.4", "title": "Platform Security (PR.PS)", "csf_function": "PR", "csf_categories": ["PR.PS"], "ferpa": [], "sox_itgc": ["Change Management", "Program Development"], "type": "ISP"},
    {"id": "ISP 4.5", "title": "Technology Infrastructure Resilience (PR.IR)", "csf_function": "PR", "csf_categories": ["PR.IR"], "ferpa": [], "sox_itgc": ["Computer Operations"], "type": "ISP"},
    {"id": "ISP 5.1", "title": "Continuous Monitoring (DE.CM)", "csf_function": "DE", "csf_categories": ["DE.CM"], "ferpa": ["§99.32"], "sox_itgc": ["Computer Operations"], "type": "ISP"},
    {"id": "ISP 5.2", "title": "Adverse Event Analysis (DE.AE)", "csf_function": "DE", "csf_categories": ["DE.AE"], "ferpa": [], "sox_itgc": [], "type": "ISP"},
    {"id": "ISP 6.1", "title": "Incident Management (RS.MA)", "csf_function": "RS", "csf_categories": ["RS.MA", "RS.CO"], "ferpa": ["§99.36", "§99.33"], "sox_itgc": ["Computer Operations"], "type": "ISP"},
    {"id": "ISP 6.2", "title": "Incident Analysis (RS.AN)", "csf_function": "RS", "csf_categories": ["RS.AN"], "ferpa": [], "sox_itgc": [], "type": "ISP"},
    {"id": "IISP 1.0", "title": "Business Continuity Policy", "csf_function": "RC", "csf_categories": ["RC.RP", "PR.IR"], "ferpa": [], "sox_itgc": ["Computer Operations"], "type": "IISP"},
    {"id": "IISP 2.0", "title": "Data Classification", "csf_function": "PR", "csf_categories": ["PR.DS", "ID.AM"], "ferpa": ["§99.3", "§99.37"], "sox_itgc": ["Access Controls"], "type": "IISP"},
    {"id": "IISP 3.0", "title": "International Travel Policy", "csf_function": "PR", "csf_categories": ["PR.DS", "PR.AA"], "ferpa": ["§99.33"], "sox_itgc": [], "type": "IISP"},
    {"id": "IISP 4.0", "title": "Acceptable Use Policy", "csf_function": "GV", "csf_categories": ["GV.PO", "PR.AT"], "ferpa": ["§99.7"], "sox_itgc": [], "type": "IISP"},
    {"id": "IISP 5.0", "title": "Artificial Intelligence Acceptable Use Policy (AI-AUP)", "csf_function": "GV", "csf_categories": ["GV.PO", "GV.RM", "PR.DS"], "ferpa": ["§99.30", "§99.31"], "sox_itgc": ["Program Development"], "type": "IISP"},
    {"id": "IISP 6.0", "title": "[RESERVED — Policy Not Yet Published]", "csf_function": "—", "csf_categories": [], "ferpa": [], "sox_itgc": [], "type": "IISP", "gap": True},
    {"id": "IISP 7.0", "title": "Vulnerability Management Policy", "csf_function": "ID", "csf_categories": ["ID.RA", "DE.CM"], "ferpa": [], "sox_itgc": ["Change Management"], "type": "IISP"},
    {"id": "IISP 8.0", "title": "Non-Employee Acceptable Use", "csf_function": "GV", "csf_categories": ["GV.PO", "PR.AA"], "ferpa": ["§99.31"], "sox_itgc": ["Access Controls"], "type": "IISP"},
    {"id": "IISP 9.0", "title": "Secure Software Development Policy", "csf_function": "PR", "csf_categories": ["PR.PS"], "ferpa": [], "sox_itgc": ["Program Development", "Change Management"], "type": "IISP"},
    {"id": "IISP 10.0", "title": "Forensics Handling of Sensitive Matters", "csf_function": "RS", "csf_categories": ["RS.AN", "RS.MA"], "ferpa": ["§99.31", "§99.36"], "sox_itgc": [], "type": "IISP"},
]

# ── ISP 1.2 Definitions ──────────────────────────────────────────────
ISP_12_DEFINITIONS = [
    ("AAR", "After Action Review — postmortem review of timeline, results, and required actions on failed or partially successful changes"),
    ("Access", "The ability to use, modify, or manipulate an information resource or gain entry to a physical area"),
    ("Agentic AI", "AI that takes initiative, understands goals, makes decisions, and carries out tasks with minimal supervision"),
    ("Artificial Intelligence", "A machine-based system that infers from input how to generate outputs such as predictions, content, recommendations, or decisions"),
    ("Break Glass Account", "Emergency account providing immediate access to critical systems when standard methods are unavailable"),
    ("Business Continuity", "Activity to ensure critical business functions remain available; encompasses Disaster Recovery"),
    ("Business Critical Applications", "Financially Significant Applications (FSAs) identified based on SOX materiality scoping"),
    ("Change", "Any alteration to state or configuration of production software or hardware under Enterprise & Learning Technologies management"),
    ("Cloud Computing", "Model for ubiquitous, on-demand network access to shared configurable computing resources"),
    ("Control", "Technical, administrative, or physical safeguard used to manage risk through preventing, detecting, or lessening threats"),
    ("Control Objective", "Target or desired condition; statement describing what is to be achieved as a result of implementing a control"),
    ("Customers", "Consumers of products without the ability to modify data identified as in the purview of privacy, financial impact, or IP"),
    ("Encryption", "Conversion of data into a form that cannot be easily understood by unauthorized people or systems"),
    ("Guideline", "Recommended practices based on industry-recognized secure practices; allows discretion unlike Standards"),
    ("Non-Human Accounts", "Accounts used for scripts, bots, APIs, or other automated processes"),
    ("Policy", "High-level statements of management intent designed to influence decisions and guide desired outcomes"),
    ("Principle of Least Privilege", "Limit access privileges for any user to those essential for assigned duties and nothing more"),
    ("Principle of Separation of Duties", "Ensure duties that can lead to fraud, abuse, or harm are separated or controlled"),
    ("Privileged User", "Someone with administrative or higher-level access beyond a typical user, performing critical system tasks"),
    ("Procedure", "Documented set of steps to perform a specific task in conformance with an applicable standard"),
    ("Risk", "Potential exposure to harm or loss, often calculated as Threat × Vulnerability × Consequence"),
    ("Service Accounts", "Accounts used by multiple people; pose security risks since tracking individual activity is difficult"),
    ("SOX", "Sarbanes-Oxley Act — US federal law setting enhanced standards for public company boards, management, and accounting"),
    ("Standard", "Mandatory requirements regarding processes, actions, and configurations designed to satisfy Control Objectives"),
    ("System Accounts", "Accounts used by operating systems or applications to perform background tasks"),
]

IISP_20_CLASSIFICATIONS = [
    ("Public", "Freely disclosed; intended for public use; minimum security controls",
     "Course catalogs, press releases, published marketing materials, regulatory filings"),
    ("Internal", "Limited to Stride employees; low sensitivity; internal controls apply",
     "Intranet, training materials, org charts, internal emails, Stride-wide policies"),
    ("Confidential", "Access based on business need-to-know; medium sensitivity; confidentiality agreements",
     "Non-directory student info, personnel records, non-public financial data, vendor contract data, source code"),
    ("Restricted / Highly Protected", "High sensitivity; strict controls; disclosure risks serious harm",
     "Passwords, private keys, government IDs (SSN, passport), financial account info, HIPAA/ADA/IDEA data, biometric data, trade secrets"),
]

# ── ASR Domains (loaded from YAML — single source of truth) ──────────
def _load_asr_domains():
    """Load asr_questions.yaml and convert to the dict format used by
    cross-reference functions and questionnaire rendering.

    YAML format:  domain["name"], q["text"], q["choices"] = [{text, risk}]
    Internal fmt: domain["domain"], q["q"],  q["choices"] = [str, ...]
    """
    yaml_path = Path(__file__).parent / "asr_questions.yaml"
    with open(yaml_path, "r", encoding="utf-8") as f:
        qb = yaml.safe_load(f)

    score_scales = {int(k): v for k, v in qb["score_scales"].items()}
    na_score = qb["na_score"]
    domains = []
    for yd in qb["domains"]:
        d = {
            "domain":      yd["name"],
            "policy_refs":  yd.get("policy_refs", []),
            "csf_refs":     yd.get("csf_refs", []),
        }
        if "ferpa_note" in yd:
            d["ferpa_note"] = yd["ferpa_note"]
        if "sox_note" in yd:
            d["sox_note"] = yd["sox_note"]

        questions = []
        for yq in yd["questions"]:
            n_c = len(yq["choices"])
            scale = score_scales[n_c]
            choice_texts = [c["text"] for c in yq["choices"]]
            choice_texts.append("N/A")
            scores = [scale[c["risk"]] for c in yq["choices"]]
            scores.append(na_score)
            questions.append({
                "q":       yq["text"],
                "choices": choice_texts,
                "weight":  yq["weight"],
                "scores":  scores,        # pre-computed per-choice scores
            })
        d["questions"] = questions
        domains.append(d)
    return domains

ASR_DOMAINS = _load_asr_domains()

# ── SRA sections ──────────────────────────────────────────────────────
SRA_SECTIONS = [
    {"num": "1", "title": "Executive Summary", "csf": ["GV.OC-01", "GV.RM-01", "GV.OV-01"]},
    {"num": "2", "title": "Trust Security Architecture Overview", "csf": ["GV.PO-01", "PR.AA-01", "PR.PS-01"]},
    {"num": "3", "title": "Personas and Trust Boundaries", "csf": ["PR.AA-01", "PR.AA-03", "ID.AM-01"]},
    {"num": "4", "title": "Identity and Access Architecture", "csf": ["PR.AA-01", "PR.AA-02", "PR.PS-01"],
     "ferpa": "FERPA §99.31 and §99.32 require identity controls that enforce role-based access aligned with educational purpose.",
     "sox": "SOX §404 AC-1 through AC-5: logical access provisioning, periodic reviews, PAM, MFA, and service account controls."},
    {"num": "5", "title": "Privileged Trust (Delinea and Entra PIM)", "csf": ["PR.AA-05", "PR.PS-02", "DE.CM-03"],
     "sox": "SOX §404 AC-3: privileged access management with JIT elevation and session recording."},
    {"num": "6", "title": "Device Trust (CrowdStrike and Defender)", "csf": ["PR.PS-03", "DE.CM-01"]},
    {"num": "7", "title": "Exposure and Posture (Tenable VM and CSPM)", "csf": ["ID.RA-01", "DE.CM-02", "PR.PS-04"]},
    {"num": "8", "title": "Application and API Trust (Sonatype, Burp Suite, Salt)", "csf": ["PR.DS-01", "PR.PS-05", "DE.CM-09"],
     "sox": "SOX §404 PD-1 through PD-2 and CM-2 through CM-6: SDLC methodology, security testing in CI/CD."},
    {"num": "9", "title": "Operational Trust (ServiceNow SecOps, CMDB, Tagging)", "csf": ["DE.CM-06", "RS.MA-01", "RC.RP-01"],
     "sox": "SOX §404 CO-1 through CO-4: job scheduling, backup/recovery, incident management."},
    {"num": "10", "title": "NIST CSF 2.0 Outcome Mapping", "csf": ["GV.OV-01", "ID.RA-01", "PR.AA-01", "DE.CM-01", "RS.MA-01", "RC.RP-01"]},
    {"num": "11", "title": "Implementation Roadmap", "csf": ["GV.SC-01", "GV.RM-02"]},
    {"num": "12", "title": "Key Performance Indicators (KPIs)", "csf": ["GV.OV-02", "DE.CM-09", "RS.MA-04"]},
    {"num": "13", "title": "Regulatory Compliance Alignment (SOX, FERPA)", "csf": ["GV.RM-03", "GV.OV-03"],
     "ferpa": "Comprehensive FERPA alignment matrix covering all applicable sections of 34 CFR Part 99.",
     "sox": "Comprehensive SOX §404 ITGC alignment matrix."},
]

DIAGRAM_MAP = {
    "1": "16-executive-summary-and-platform-stack.png",
    "2a": "01-trust-security-architecture-overview.png",
    "2b": "02-trust-logical-architecture-components.png",
    "3": "03-persona-trust-boundaries.png",
    "4a": "04-staff-authentication-flow.png",
    "4b": "05-staff-conditional-access-policy-matrix.png",
    "4c": "06-student-authentication-flow.png",
    "4d": "07-customer-authentication-flow.png",
    "4e": "08-identity-and-access-architecture.png",
    "5": "09-privileged-trust-architecture.png",
    "6": "10-device-trust-architecture.png",
    "7": "11-exposure-and-posture-architecture.png",
    "8": "12-application-and-api-trust-architecture.png",
    "9": "13-operational-trust-architecture.png",
    "10": "14-nist-csf-outcome-mapping.png",
    "11": "15-implementation-roadmap.png",
    "12": "17-kpi-framework.png",
    "13": "18-regulatory-compliance-alignment.png",
}

SEC_DIAGRAM_KEYS = {
    "1": ["1"], "2": ["2a", "2b"], "3": ["3"],
    "4": ["4a", "4b", "4c", "4d", "4e"], "5": ["5"],
    "6": ["6"], "7": ["7"], "8": ["8"], "9": ["9"],
    "10": ["10"], "11": ["11"], "12": ["12"], "13": ["13"],
}

# ── SRA Inter-Section Dependencies (architectural) ───────────────────
# Maps SRA section → list of related SRA sections it builds upon or informs
SRA_SECTION_DEPS = {
    "1": ["2", "10", "11"],       # Executive Summary references overview, mapping, roadmap
    "2": ["3", "4", "5", "6", "7", "8", "9"],  # Overview underpins all operational sections
    "3": ["4"],                    # Personas feed identity architecture
    "4": ["5", "3"],              # IAM depends on privileged trust and personas
    "5": ["4"],                    # Privileged trust depends on IAM
    "6": ["7", "9"],              # Device trust feeds posture management and ops
    "7": ["6", "8"],              # Exposure management covers devices and app security
    "8": ["7", "4"],              # App trust uses vulnerability & IAM controls
    "9": ["6", "7", "8", "5"],    # Ops trust spans device, exposure, app, privileged
    "10": ["1", "2", "3", "4", "5", "6", "7", "8", "9"],  # CSF mapping covers all
    "11": ["10", "12"],            # Roadmap informed by mapping and KPIs
    "12": ["10", "9"],             # KPIs measure outcomes and ops
    "13": ["4", "5", "8", "9"],    # Regulatory alignment focuses on SOX/FERPA-relevant sections
}


def _csf_category(subcat):
    """Extract CSF category from a subcategory code, e.g. 'PR.AA-01' → 'PR.AA'."""
    return re.sub(r'-\d+$', '', subcat)


def _build_sra_to_policies():
    """Map SRA section# → list of policy IDs via shared CSF categories."""
    result = {}
    for sec in SRA_SECTIONS:
        sec_cats = {_csf_category(c) for c in sec["csf"]}
        matched = []
        for pol in POLICY_INVENTORY:
            if pol.get("gap"):
                continue
            pol_cats = set(pol["csf_categories"])
            if sec_cats & pol_cats:
                matched.append(pol["id"])
        result[sec["num"]] = matched
    return result


def _build_sra_to_asr():
    """Map SRA section# → list of (domain_name, question_numbers) via shared CSF categories."""
    result = {}
    for sec in SRA_SECTIONS:
        sec_cats = {_csf_category(c) for c in sec["csf"]}
        matched = []
        q_offset = 0
        for domain in ASR_DOMAINS:
            dom_cats = set(domain["csf_refs"])
            n_q = len(domain["questions"])
            if sec_cats & dom_cats:
                q_range = f"Q{q_offset + 1}–Q{q_offset + n_q}"
                matched.append((domain["domain"], q_range))
            q_offset += n_q
        result[sec["num"]] = matched
    return result


def _build_policy_to_asr():
    """Map policy ID → list of (domain_name, question_numbers) via direct policy_refs."""
    result = {}
    q_offset = 0
    offsets = []
    for domain in ASR_DOMAINS:
        offsets.append((domain, q_offset))
        q_offset += len(domain["questions"])

    for pol in POLICY_INVENTORY:
        if pol.get("gap"):
            continue
        pid = pol["id"]
        matched = []
        for domain, off in offsets:
            if pid in domain["policy_refs"]:
                n_q = len(domain["questions"])
                q_range = f"Q{off + 1}–Q{off + n_q}"
                matched.append((domain["domain"], q_range))
        result[pid] = matched
    return result


def _build_policy_to_sra():
    """Map policy ID → list of SRA section nums via shared CSF categories."""
    result = {}
    for pol in POLICY_INVENTORY:
        if pol.get("gap"):
            continue
        pol_cats = set(pol["csf_categories"])
        matched = []
        for sec in SRA_SECTIONS:
            sec_cats = {_csf_category(c) for c in sec["csf"]}
            if pol_cats & sec_cats:
                matched.append(sec["num"])
        result[pol["id"]] = matched
    return result


# Precompute all cross-reference maps
SRA_TO_POLICIES = _build_sra_to_policies()
SRA_TO_ASR = _build_sra_to_asr()
POLICY_TO_ASR = _build_policy_to_asr()
POLICY_TO_SRA = _build_policy_to_sra()

# ── SRA Acronyms Table ────────────────────────────────────────────────
SRA_ACRONYMS = [
    ("CA", "Conditional Access — policy-based access control in Microsoft Entra"),
    ("CIAM", "Customer Identity and Access Management — IAM for external/customer-facing authentication (Okta CIC)"),
    ("CMDB", "Configuration Management Database — central repository of IT asset records (ServiceNow CMDB)"),
    ("CSF", "Cybersecurity Framework — NIST framework (v2.0) across Govern, Identify, Protect, Detect, Respond, Recover"),
    ("CSPM", "Cloud Security Posture Management — continuous monitoring of cloud configurations (Tenable Cloud Security)"),
    ("EDR", "Endpoint Detection and Response — real-time endpoint monitoring and threat response (CrowdStrike Falcon)"),
    ("FERPA", "Family Educational Rights and Privacy Act — 34 CFR Part 99, protecting student education records"),
    ("FIDO2", "Fast Identity Online 2 — passwordless authentication standard using hardware security keys"),
    ("IAM", "Identity and Access Management — policies and technologies for right individuals/right resources/right time"),
    ("IGA", "Identity Governance and Administration — lifecycle management, access certification, and compliance reporting"),
    ("ITGC", "IT General Controls — SOX §404 controls over access, change management, computer operations"),
    ("JEA", "Just-Enough Administration — restricting privileged sessions to minimum required commands"),
    ("JIT", "Just-in-Time Access — time-bound privilege elevation that expires automatically (Entra PIM, Delinea)"),
    ("MDM", "Mobile Device Management — device enrollment, policy enforcement, remote wipe"),
    ("MFA", "Multi-Factor Authentication — requiring two or more verification factors for access"),
    ("MTTR", "Mean Time to Remediate/Respond — average time between detection and resolution"),
    ("NIST", "National Institute of Standards and Technology — U.S. federal cybersecurity standards"),
    ("PAM", "Privileged Access Management — discovering, managing, auditing privileged accounts (Delinea)"),
    ("PIM", "Privileged Identity Management — role-based JIT elevation with approval workflows (Entra PIM)"),
    ("RBAC", "Role-Based Access Control — access permissions assigned by organizational role"),
    ("SecOps", "Security Operations — monitoring, detecting, and responding to security events"),
    ("SLA", "Service Level Agreement — formal commitment on availability, response, or remediation timelines"),
    ("SOX", "Sarbanes-Oxley Act — U.S. legislation requiring internal controls over financial reporting (§404)"),
    ("SRA", "Security Reference Architecture — this document; the enterprise trust and security architecture blueprint"),
    ("TOGAF", "The Open Group Architecture Framework — enterprise architecture methodology"),
    ("VM", "Vulnerability Management — identifying, classifying, prioritizing, remediating vulnerabilities (Tenable)"),
    ("WHfB", "Windows Hello for Business — passwordless, certificate-based authentication using biometrics or PIN"),
    ("ZT", "Zero Trust — continuous verification of every user, device, and transaction (NIST SP 800-207)"),
]

# ── CSF-to-SRA Mapping (for SRA Appendix C) ──────────────────────────
CSF_SRA_MAP = [
    ("GV.OC-01", "Organizational context for cybersecurity risk management", "§1 Executive Summary", "Implemented"),
    ("GV.RM-01", "Risk management objectives established and communicated", "§1 Executive Summary", "Implemented"),
    ("GV.RM-02", "Risk appetite and tolerance statements defined", "§11 Implementation Roadmap", "In Progress"),
    ("GV.RM-03", "Risk management activities integrated across enterprise", "§13 Regulatory Compliance", "Implemented"),
    ("GV.OV-01", "Cybersecurity strategy outcomes assessed", "§1, §10 NIST CSF Mapping", "Implemented"),
    ("GV.OV-02", "Cybersecurity performance evaluated against KPIs", "§12 KPIs", "Implemented"),
    ("GV.PO-01", "Cybersecurity policy established and communicated", "§2 Architecture Overview", "Implemented"),
    ("ID.AM-01", "Asset inventories maintained", "§3 Personas and Trust Boundaries", "Implemented"),
    ("ID.RA-01", "Vulnerabilities identified, validated, and recorded", "§7 Exposure and Posture, §10", "Implemented"),
    ("PR.AA-01", "Identities and credentials managed", "§4 Identity and Access", "Implemented"),
    ("PR.AA-05", "Least privilege and separation of duties enforced", "§5 Privileged Trust", "Implemented"),
    ("PR.DS-01", "Data-at-rest and data-in-transit protected", "§8 Application and API Trust", "Implemented"),
    ("PR.PS-01", "Configuration baselines maintained", "§2, §4 Architecture Overview", "Implemented"),
    ("PR.PS-03", "Hardware maintained and replaced per policy", "§6 Device Trust", "Implemented"),
    ("DE.CM-01", "Networks monitored for anomalies and attacks", "§6 Device Trust", "Implemented"),
    ("DE.CM-03", "Personnel activity monitored for anomalies", "§5 Privileged Trust", "Implemented"),
    ("RS.MA-01", "Incident response plan executed", "§9 Operational Trust", "Implemented"),
    ("RC.RP-01", "Recovery plan executed during/after incident", "§9 Operational Trust", "Implemented"),
]

EVIDENCE_MAP = [
    ("GV.OC-01", "SRA document, governance charter", "Security Architecture"),
    ("GV.RM-01", "Risk register, risk appetite statement", "GRC / Risk Management"),
    ("GV.OV-01", "CSF outcome mapping (§10), annual review records", "Security Architecture"),
    ("GV.OV-02", "KPI dashboard (§12), quarterly metrics reports", "Security Operations"),
    ("GV.PO-01", "ISP 1.0–12.0 policy library, IISP 2.0", "Information Security"),
    ("ID.AM-01", "ServiceNow CMDB asset inventory, Entra directory", "IT Operations"),
    ("ID.RA-01", "Tenable scan reports, risk assessment records", "Vulnerability Management"),
    ("PR.AA-01", "Entra ID configurations, Okta CIC tenant configs", "Identity Engineering"),
    ("PR.AA-05", "Delinea session recordings, PIM activation logs", "Privileged Access Mgmt"),
    ("PR.DS-01", "TLS configurations, encryption-at-rest settings", "Application Security"),
    ("PR.PS-01", "CIS benchmark scan results, baseline configs", "Infrastructure Security"),
    ("DE.CM-01", "CrowdStrike detection logs, network alert reports", "Security Operations"),
    ("DE.CM-03", "PIM audit logs, privileged session recordings", "Security Operations"),
    ("RS.MA-01", "ServiceNow incident records, runbook executions", "Incident Response"),
    ("RC.RP-01", "Recovery test results, backup verification logs", "Business Continuity"),
]


# ── Handcrafted Policy Summaries (bullet lists) ──────────────────────
POLICY_SUMMARIES = {
    "ISP 1.0": [
        "Establishes Stride's enterprise information security program under executive authority.",
        "All business units must comply; the CISO is the designated accountable officer.",
        "Defines the governance structure, risk management approach, and policy hierarchy.",
        "Requires annual review and board-level reporting on cybersecurity posture.",
        "Mandates alignment with NIST CSF 2.0, FERPA, and SOX §404 requirements.",
    ],
    "ISP 1.1": [
        "Serves as the master index for all Stride IT Security Policies and Standards.",
        "Enumerates ISP 1.0 through ISP 6.2 and IISP 1.0 through IISP 10.0.",
        "Establishes the policy numbering convention and document hierarchy.",
        "Applies to all Stride-owned technology and business operations.",
    ],
    "ISP 1.2": [
        "Defines key terms used across the Stride information security policy framework.",
        "Covers access, controls, encryption, risk, privileged users, and policy types.",
        "Distinguishes Policies (management intent), Standards (mandatory), and Guidelines (recommended).",
        "Includes AI-specific terminology: Agentic AI and Artificial Intelligence.",
        "Defines account types: service, system, non-human, and break-glass accounts.",
    ],
    "ISP 2.1": [
        "Establishes the organizational context for cybersecurity risk management decisions.",
        "Aligns cybersecurity strategies with Stride's educational mission and business objectives.",
        "Requires documentation of internal and external factors shaping security priorities.",
        "Mandates understanding of stakeholder expectations, legal obligations, and the risk environment.",
    ],
    "ISP 2.2": [
        "Defines Stride's risk management strategy, including risk appetite and tolerance levels.",
        "Establishes a structured approach to identifying, assessing, and treating cybersecurity risks.",
        "Requires risk management to be integrated across all enterprise processes.",
        "Mandates periodic review of risk tolerance statements and alignment with business goals.",
    ],
    "ISP 2.3": [
        "Defines roles, responsibilities, and authorities for cybersecurity governance.",
        "Assigns security functions to the CISO, IT leadership, and business unit owners.",
        "Requires separation of duties and least privilege for high-risk operations.",
        "Mandates RACI matrices for key security processes and incident response.",
        "Establishes accountability for policy compliance at all organizational levels.",
    ],
    "ISP 2.4": [
        "Establishes the framework for creating, communicating, and enforcing security policies.",
        "Requires documented processes and procedures for all security-relevant operations.",
        "Mandates annual notification of FERPA rights to parents and eligible students.",
        "Defines the policy lifecycle: development, review, approval, distribution, and retirement.",
    ],
    "ISP 2.5": [
        "Defines oversight mechanisms for risk management effectiveness.",
        "Requires periodic assessment of cybersecurity strategy outcomes.",
        "Establishes executive and board-level security posture reporting requirements.",
        "Mandates coordination with regulatory enforcement bodies (Family Policy Compliance Office).",
    ],
    "ISP 2.6": [
        "Defines standards for managing cybersecurity risks in the supply chain.",
        "Requires vendor security assessments and contractual security obligations.",
        "Mandates controls on data sharing with third parties, including redisclosure limitations.",
        "Applies to all outsourced services, SaaS platforms, and contractor relationships.",
        "Requires exit strategies and data portability plans for critical vendor dependencies.",
    ],
    "ISP 3.1": [
        "Defines standards for maintaining a comprehensive inventory of IT assets.",
        "Requires all hardware, software, and data assets to be cataloged in the CMDB.",
        "Mandates classification of assets by business criticality and data sensitivity.",
        "Establishes lifecycle management from acquisition through disposal.",
    ],
    "ISP 3.2": [
        "Defines standards for conducting cybersecurity risk assessments.",
        "Requires identification, validation, and recording of vulnerabilities enterprise-wide.",
        "Mandates risk assessment for health/safety emergencies per FERPA §99.36.",
        "Establishes risk scoring methodology and remediation prioritization criteria.",
        "Requires assessment results to be integrated into the enterprise risk register.",
    ],
    "ISP 3.3": [
        "Defines standards for continuous improvement of security response capabilities.",
        "Requires lessons learned from incidents, assessments, and exercises to drive program changes.",
        "Mandates after-action reviews (AARs) for significant security events.",
        "Establishes metrics for measuring improvement in detection and response effectiveness.",
    ],
    "ISP 4.1": [
        "Defines standards for identity management, authentication, and access control.",
        "Requires MFA for all administrative and FERPA-sensitive system access.",
        "Mandates role-based access control aligned with least privilege principles.",
        "Establishes user access review cycles: quarterly (privileged), semi-annual (standard).",
        "Requires audit logging of authentication events and access changes per FERPA §99.32.",
    ],
    "ISP 4.2": [
        "Defines standards for security awareness and training programs.",
        "Requires annual security awareness training for all workforce members.",
        "Mandates role-specific training for privileged users and IT administrators.",
        "Includes FERPA awareness training covering student data privacy obligations.",
        "Requires phishing simulation exercises and training effectiveness measurement.",
    ],
    "ISP 4.3": [
        "Defines standards for protecting data at rest and in transit.",
        "Requires AES-256 or equivalent encryption for sensitive and confidential data.",
        "Mandates TLS 1.2+ for all data in transit.",
        "Establishes data handling procedures aligned with FERPA consent and disclosure rules.",
        "Requires data loss prevention (DLP) controls for PII and education records.",
    ],
    "ISP 4.4": [
        "Defines standards for platform security, configuration management, and patching.",
        "Requires CIS benchmark-aligned hardening for all server and endpoint platforms.",
        "Mandates segregation of development, test, and production environments.",
        "Establishes change management processes aligned with SOX §404 ITGC requirements.",
        "Requires automated configuration compliance monitoring and drift detection.",
    ],
    "ISP 4.5": [
        "Defines standards for technology infrastructure resilience.",
        "Requires high-availability architectures for business-critical systems.",
        "Mandates redundancy, failover, and disaster recovery capabilities.",
        "Establishes infrastructure monitoring and capacity planning requirements.",
    ],
    "ISP 5.1": [
        "Defines standards for continuous security monitoring across the enterprise.",
        "Requires SIEM integration with real-time alerting for all critical systems.",
        "Mandates network, endpoint, and application-level monitoring.",
        "Establishes log retention requirements supporting audit and forensic needs.",
        "Requires monitoring of disclosure records per FERPA §99.32.",
    ],
    "ISP 5.2": [
        "Defines standards for analyzing adverse events and anomalies.",
        "Requires correlation of security events across multiple data sources.",
        "Mandates triage procedures and severity classification for detected events.",
        "Establishes escalation thresholds and notification requirements.",
    ],
    "ISP 6.1": [
        "Defines standards for incident response planning and execution.",
        "Requires a documented IRP with application-specific runbooks.",
        "Mandates defined escalation paths, notification responsibilities, and communication plans.",
        "Establishes incident severity levels and response time SLAs.",
        "Requires coordination with legal counsel for FERPA breach notifications.",
    ],
    "ISP 6.2": [
        "Defines standards for incident analysis and root-cause investigation.",
        "Requires systematic analysis to determine scope, impact, and root cause.",
        "Mandates evidence preservation and chain-of-custody procedures.",
        "Establishes post-incident review processes to drive program improvements.",
    ],
    "IISP 1.0": [
        "Establishes business continuity requirements for critical operations.",
        "Requires business impact analysis (BIA) to identify essential functions and dependencies.",
        "Mandates documented recovery procedures with defined RTO and RPO targets.",
        "Requires annual business continuity plan testing and tabletop exercises.",
        "Covers disaster recovery, communication plans, and workforce continuity.",
    ],
    "IISP 2.0": [
        "Classifies institutional data into four tiers based on sensitivity and regulatory impact.",
        "Public: freely disclosed; minimum controls (e.g., press releases, course catalogs).",
        "Internal: limited to employees; low sensitivity (e.g., org charts, training materials).",
        "Confidential: need-to-know; medium sensitivity (e.g., student records, source code).",
        "Restricted: strictest controls (e.g., SSNs, private keys, biometric data).",
    ],
    "IISP 3.0": [
        "Establishes security requirements for international travel with company assets.",
        "Requires pre-travel risk assessment and approval for devices carrying sensitive data.",
        "Mandates use of loaner/clean devices for travel to high-risk regions.",
        "Requires post-travel device inspection and credential rotation.",
    ],
    "IISP 4.0": [
        "Defines acceptable and responsible use of Stride's information systems and networks.",
        "Applies to all employees, contractors, and authorized users of Stride technology.",
        "Prohibits unauthorized access, data exfiltration, and misuse that compromises security.",
        "Requires compliance with FERPA data handling obligations.",
        "Establishes monitoring rights and disciplinary consequences for violations.",
    ],
    "IISP 5.0": [
        "Governs the use of Artificial Intelligence technology in Stride computing environments.",
        "Requires risk assessment for AI tools that process student data or inform decisions.",
        "Mandates human review of AI outputs before use in production or decision-making.",
        "Prohibits input of restricted/confidential data into unapproved AI systems.",
        "Requires phased deployment (crawl/walk/run) and legal review of AI vendor terms.",
    ],
    "IISP 7.0": [
        "Establishes a comprehensive vulnerability management program.",
        "Requires regular vulnerability scanning (authenticated, weekly for critical systems).",
        "Defines remediation SLAs: 7 days critical, 30 days high, 90 days medium.",
        "Mandates integration with change management for security patches.",
        "Requires vulnerability metrics reporting to CISO and governance bodies.",
    ],
    "IISP 8.0": [
        "Defines minimum security and privacy requirements for non-employee system access.",
        "Applies to contractors, consultants, vendors, and other third-party personnel.",
        "Requires signed acceptable use agreements before granting any system access.",
        "Mandates access reviews and prompt deprovisioning upon engagement termination.",
        "Prohibits access to education records unless authorized under FERPA exceptions.",
    ],
    "IISP 9.0": [
        "Defines mandatory requirements for secure software development and acquisition.",
        "Requires SAST, DAST, and SCA testing integrated into the CI/CD pipeline.",
        "Mandates peer code review and security gates at each SDLC phase.",
        "Establishes secure coding standards and developer security training requirements.",
        "Requires environment segregation per SOX §404 ITGC.",
    ],
    "IISP 10.0": [
        "Establishes requirements for forensic handling of digital devices and data.",
        "Applies to legal, HR, regulatory, and law enforcement investigations.",
        "Requires chain-of-custody documentation and evidence integrity controls.",
        "Mandates qualified forensic examiners and approved tools for evidence collection.",
        "Prohibits modification of potential evidence prior to formal preservation.",
    ],
}


def get_policy_summary_bullets(policy_id):
    """Return a list of bullet-point summary strings for a policy."""
    return POLICY_SUMMARIES.get(policy_id, ["Summary not available."])


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Policy Cross Reference V7
# ══════════════════════════════════════════════════════════════════════

def build_policy_cross_reference_docx():
    print("Building DOCX: Stride Policy Cross Reference V7...")
    doc = new_doc()
    make_landscape(doc)

    # Title page
    add_title_page(doc, "Stride Policy Cross Reference", "Version 7.0",
                   extra_lines=[
                       "Regulatory Cross-Reference Frameworks:",
                       "NIST Cybersecurity Framework 2.0",
                       "FERPA (34 CFR Part 99)",
                       "SOX Section 404 — IT General Controls",
                       "",
                       "Note: IISP 6.0 is reserved and not yet published.",
                   ])

    add_header_footer(doc, "Stride, Inc. — Policy Cross Reference V7")

    # Table of Contents
    add_table_of_contents(doc)

    # Section 1: Policy Inventory Overview
    doc.add_heading("1. Policy Inventory Overview", level=1)
    doc.add_paragraph(
        "The Stride information security program comprises 21 Information Security Policies (ISPs) "
        "aligned to NIST CSF 2.0 functions and 10 Integrated Information Security Policies (IISPs) "
        "addressing cross-cutting operational domains. Together, these 31 policy documents "
        "(30 published + 1 reserved) establish the governance foundation for cybersecurity risk management."
    )

    # CSF coverage list
    coverage = [
        ("GOVERN (GV):", "ISP 1.0–1.2, ISP 2.1–2.6, IISP 4.0, 5.0, 8.0"),
        ("IDENTIFY (ID):", "ISP 3.1–3.3, IISP 7.0"),
        ("PROTECT (PR):", "ISP 4.1–4.5, IISP 2.0, 3.0, 9.0"),
        ("DETECT (DE):", "ISP 5.1–5.2"),
        ("RESPOND (RS):", "ISP 6.1–6.2, IISP 10.0"),
        ("RECOVER (RC):", "IISP 1.0"),
    ]
    for func, pols in coverage:
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_run(p, func, bold=True)
        p.add_run(f" {pols}")

    p = doc.add_paragraph()
    add_formatted_run(p, "Gap: ", bold=True, color=GAP_RED)
    add_formatted_run(p, "IISP 6.0 is reserved and not yet published.", color=GAP_RED)

    doc.add_page_break()

    # Section 2: Per-Policy Detail
    doc.add_heading("2. Per-Policy Cross-Reference Detail", level=1)

    for pol in POLICY_INVENTORY:
        pid = pol["id"]
        title = pol["title"]
        is_gap = pol.get("gap", False)

        heading_text = f"STRIDE {pid} — {title}"
        h = doc.add_heading(heading_text, level=2)
        if is_gap:
            for run in h.runs:
                run.font.color.rgb = GAP_RED

        # Summary (bullet points)
        if is_gap:
            p = doc.add_paragraph()
            add_formatted_run(p, "This policy number is reserved but has not yet been published. "
                              "This represents a gap in the current policy framework.", italic=True, color=GAP_RED)
        else:
            bullets = get_policy_summary_bullets(pid)
            p = doc.add_paragraph()
            add_formatted_run(p, "Summary:", bold=True)
            for bullet in bullets:
                doc.add_paragraph(bullet, style="List Bullet")

        # ISP 1.2: Merged definitions + cross-reference table
        if pid == "ISP 1.2":
            p = doc.add_paragraph()
            add_formatted_run(p, "Key Definitions and Cross-Reference:", bold=True)
            # Build one combined table: Term | Definition
            tbl = doc.add_table(rows=1, cols=2)
            tbl.rows[0].cells[0].text = "Term"
            tbl.rows[0].cells[1].text = "Definition"
            for term, defn in ISP_12_DEFINITIONS:
                row = tbl.add_row()
                row.cells[0].text = term
                row.cells[1].text = defn
            style_table(tbl)
            # Still add the cross-ref table below

        # IISP 2.0 Classification enrichment
        if pid == "IISP 2.0":
            p = doc.add_paragraph()
            add_formatted_run(p, "Data Classification Levels:", bold=True)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.rows[0].cells[0].text = "Level"
            tbl.rows[0].cells[1].text = "Description & Controls"
            tbl.rows[0].cells[2].text = "Examples"
            for level, desc, examples in IISP_20_CLASSIFICATIONS:
                row = tbl.add_row()
                row.cells[0].text = level
                row.cells[1].text = desc
                row.cells[2].text = examples
            style_table(tbl)

        # Cross-reference table — base rows + optional SRA/ASR cross-refs
        # Each entry is (label, value) where value is str or list[str].
        # Lists render one item per line via cell.add_paragraph().
        xref_rows = [
            ("NIST CSF 2.0 Function", f"{pol['csf_function']} — {NIST_CSF_FUNCTIONS.get(pol['csf_function'], '—')}"),
            ("NIST CSF 2.0 Categories", ", ".join(pol["csf_categories"]) if pol["csf_categories"] else "—"),
            ("FERPA (34 CFR Part 99)", ", ".join(pol["ferpa"]) if pol["ferpa"] else "—"),
            ("SOX §404 ITGC Domains", ", ".join(pol["sox_itgc"]) if pol["sox_itgc"] else "—"),
        ]
        if not is_gap:
            sra_nums = POLICY_TO_SRA.get(pid, [])
            if sra_nums:
                sra_labels = []
                for sn in sra_nums:
                    for s in SRA_SECTIONS:
                        if s["num"] == sn:
                            sra_labels.append(f"§{sn} {s['title']}")
                            break
                xref_rows.append(("SRA Sections", sra_labels))
            asr_refs = POLICY_TO_ASR.get(pid, [])
            if asr_refs:
                labels = [f"{dom} ({qr})" for dom, qr in asr_refs]
                xref_rows.append(("ASR Questionnaire Coverage", labels))

        tbl = doc.add_table(rows=1 + len(xref_rows), cols=2)
        tbl.rows[0].cells[0].text = "Framework"
        tbl.rows[0].cells[1].text = "Alignment"
        for ri, (label, value) in enumerate(xref_rows, start=1):
            tbl.rows[ri].cells[0].text = label
            if isinstance(value, list):
                # First item goes in the default paragraph, rest via add_paragraph
                cell = tbl.rows[ri].cells[1]
                cell.text = value[0]
                for item in value[1:]:
                    cell.add_paragraph(item)
            else:
                tbl.rows[ri].cells[1].text = value

        style_table(tbl)

    doc.add_page_break()

    # Appendix A: FERPA Reverse Map
    doc.add_heading("APPENDIX A — FERPA (34 CFR Part 99) Reverse Mapping", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.rows[0].cells[0].text = "Section"
    tbl.rows[0].cells[1].text = "Description"
    tbl.rows[0].cells[2].text = "Stride Policies"
    for section_id, section_desc in FERPA_SECTIONS.items():
        matching = [p["id"] for p in POLICY_INVENTORY if section_id in p["ferpa"]]
        row = tbl.add_row()
        row.cells[0].text = section_id
        row.cells[1].text = section_desc
        row.cells[2].text = ", ".join(matching) if matching else "No direct mapping"
    style_table(tbl)

    doc.add_page_break()

    # Appendix B: SOX §404 Reverse Map — Tabular
    doc.add_heading("APPENDIX B — SOX §404 ITGC Reverse Mapping", level=1)
    doc.add_paragraph(
        "Maps SOX §404 ITGC domains to their representative controls and the Stride policies "
        "that provide coverage. Empty cells in the Stride Policies column indicate potential coverage gaps."
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "SOX §404 Domain"
    tbl.rows[0].cells[1].text = "Control ID / Description"
    tbl.rows[0].cells[2].text = "Stride Policies"
    tbl.rows[0].cells[3].text = "NIST CSF Alignment"

    # Map SOX domains to CSF functions for the alignment column
    sox_csf_map = {
        "Access Controls": "PR.AA, GV.RR",
        "Change Management": "PR.PS, GV.PO",
        "Computer Operations": "DE.CM, RS.MA, RC.RP",
        "Program Development": "PR.PS, ID.RA",
    }
    for domain, domain_data in SOX_404_ITGCS.items():
        matching = [p["id"] for p in POLICY_INVENTORY if domain in p["sox_itgc"]]
        policies_str = ", ".join(matching) if matching else "—"
        csf_str = sox_csf_map.get(domain, "—")
        for ctrl in domain_data["controls"]:
            row = tbl.add_row()
            row.cells[0].text = domain
            row.cells[1].text = ctrl
            row.cells[2].text = policies_str
            row.cells[3].text = csf_str
    style_table(tbl)

    doc.add_page_break()

    # Appendix C: ASR Coverage Gap Analysis
    doc.add_heading("APPENDIX C — ASR Questionnaire Coverage Gap Analysis", level=1)
    doc.add_paragraph(
        "The Application Security Review (ASR) questionnaire identifies control expectations across seven "
        "domains. This appendix maps ASR questions to existing Stride policies and highlights topics where "
        "current policy coverage is implicit, indirect, or absent — representing areas for future policy development."
    )

    # Identify gaps: questions whose topics go beyond the referenced policies
    ASR_GAP_ITEMS = [
        {
            "domain": "Governance and Program Management",
            "question": "What is the business necessity classification for this application?",
            "gap_type": "Implicit",
            "current_coverage": "ISP 3.1 (Asset Management) identifies assets but does not define a formal business-necessity classification (KTLO / SHOULD / MAY).",
            "recommendation": "Formalize a business-necessity taxonomy as a standard under ISP 3.1 or a new annex.",
        },
        {
            "domain": "Governance and Program Management",
            "question": "Are productivity improvement estimates or benchmarks documented?",
            "gap_type": "Absent",
            "current_coverage": "No published policy addresses productivity benchmarking for technology adoption.",
            "recommendation": "Consider adding quantitative benefit documentation requirements to the enterprise architecture governance process.",
        },
        {
            "domain": "Governance and Program Management",
            "question": "Has legal counsel reviewed this application's terms, data handling, and IP implications?",
            "gap_type": "Implicit",
            "current_coverage": "IISP 5.0 (AI AUP) references legal review for AI tools. No general-purpose legal review policy exists for all application types.",
            "recommendation": "Establish a legal/IP review gate within the application onboarding process; codify in ISP 2.4 or new standard.",
        },
        {
            "domain": "Governance and Program Management",
            "question": "Has the application been reviewed through the enterprise architecture governance process?",
            "gap_type": "Implicit",
            "current_coverage": "ISP 2.5 (Oversight) addresses governance broadly. No dedicated enterprise architecture review policy or standard exists.",
            "recommendation": "Publish an EA governance review standard linked to ISP 2.5, including defined approval gates.",
        },
        {
            "domain": "Secure Development and Change Management",
            "question": "Does the application require an out-of-band update mechanism?",
            "gap_type": "Partial",
            "current_coverage": "ISP 4.4 (Platform Security) covers patching. Out-of-band update channels (CLI, vendor-pushed) lack specific controls.",
            "recommendation": "Add out-of-band update requirements to ISP 4.4 or IISP 9.0 covering approval, validation, and rollback.",
        },
        {
            "domain": "Secure Development and Change Management",
            "question": "Are human review and approval required before application outputs are used?",
            "gap_type": "Partial",
            "current_coverage": "IISP 5.0 requires human review for AI-generated outputs. No equivalent policy for non-AI application outputs that drive decisions.",
            "recommendation": "Extend human-in-the-loop requirements beyond AI to any application generating operational or financial decisions.",
        },
        {
            "domain": "Vulnerability and Threat Management",
            "question": "Is application usage monitored with defined KPIs?",
            "gap_type": "Implicit",
            "current_coverage": "ISP 5.1 (Continuous Monitoring) covers security event monitoring. Application usage telemetry and KPI reporting are not formalized.",
            "recommendation": "Define usage monitoring and KPI requirements as a standard under ISP 5.1.",
        },
        {
            "domain": "Third-Party and Supply Chain Risk",
            "question": "Is there a phased deployment plan with maturity-gated control enhancements?",
            "gap_type": "Implicit",
            "current_coverage": "IISP 5.0 references a crawl/walk/run approach for AI. No general phased deployment framework for other application types.",
            "recommendation": "Generalize the phased deployment model from IISP 5.0 into ISP 2.6 or a new standard applicable to all third-party/SaaS onboarding.",
        },
    ]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.rows[0].cells[0].text = "Domain"
    tbl.rows[0].cells[1].text = "ASR Question"
    tbl.rows[0].cells[2].text = "Gap Type"
    tbl.rows[0].cells[3].text = "Current Policy Coverage"
    tbl.rows[0].cells[4].text = "Recommendation"

    gap_colors = {"Absent": "FFCDD2", "Implicit": "FFF9C4", "Partial": "FFE0B2"}
    for item in ASR_GAP_ITEMS:
        row = tbl.add_row()
        row.cells[0].text = item["domain"]
        row.cells[1].text = item["question"]
        row.cells[2].text = item["gap_type"]
        row.cells[3].text = item["current_coverage"]
        row.cells[4].text = item["recommendation"]
        bg = gap_colors.get(item["gap_type"])
        if bg:
            set_cell_shading(row.cells[2], bg)
    style_table(tbl)

    p = doc.add_paragraph()
    add_formatted_run(p, "Gap Types:", bold=True)
    doc.add_paragraph("Absent \u2014 no published policy addresses this topic.", style="List Bullet")
    doc.add_paragraph("Implicit \u2014 existing policy partially relevant but not explicit.", style="List Bullet")
    doc.add_paragraph("Partial \u2014 dedicated policy exists but does not fully address the ASR requirement.", style="List Bullet")

    ensure_copyable(doc)
    out_path = OUTPUT / "Stride_Policy_Cross_Reference_V7.docx"
    doc.save(str(out_path))
    print(f"  SUCCESS: {out_path} ({out_path.stat().st_size:,} bytes)")
    return out_path


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 2: SRA V8
# ══════════════════════════════════════════════════════════════════════

def build_sra_v8_docx():
    print("Building DOCX: Stride Security Reference Architecture V8...")
    doc = new_doc()
    make_landscape(doc)

    add_title_page(doc, "Stride Security Reference Architecture", "Version 8.0",
                   extra_lines=[
                       "Enterprise Trust Security Reference Architecture",
                       "Incorporating:",
                       "22 D2 Architecture Diagrams (Landscape-Optimized)",
                       "Per-Section NIST CSF 2.0 Cross-References",
                       "FERPA and SOX §404 Regulatory Alignment",
                       "Cross-References to Policies and ASR Questionnaire",
                       "",
                       "Supersedes: SRA V7 (March 3, 2026)",
                   ])

    add_header_footer(doc, "Stride Security Reference Architecture V8")

    # Table of Contents, then Table of Figures (Word populates on first update)
    add_table_of_contents(doc)
    add_table_of_figures(doc)

    # Load SRA content
    sra_enhanced = workpapers_data.get("SRA Enhanced", {})
    sra_v6 = workpapers_data.get("SRA v6", {})
    sra_source = sra_enhanced if sra_enhanced.get("chars", 0) > sra_v6.get("chars", 0) else sra_v6
    sra_paras = sra_source.get("structured", [])
    sra_anchored = workpapers_data.get("Stride_SRA_Consolidated_Anchored_D2", {})
    sra_anchored_paras = sra_anchored.get("structured", [])

    _STRIP_RE = re.compile(
        r'^\[DIAGRAM PLACEHOLDER\]'
        r'|^Control\s*/\s*Policy\s*/\s*Compliance'
        r'|^Tightened Augmentation'
        r'|^Appendix\s+[A-Z]'
        r'|^See\s+SRA\s+Enhanced'
        r'|^Evidence\s+sources\s+are\s+referenced',
        re.IGNORECASE,
    )

    def extract_section_content(section_num, paras):
        in_section = False
        raw_lines = []
        section_pattern = re.compile(rf'^{re.escape(section_num)}[\.\s]')
        next_section = str(int(section_num) + 1) if section_num.isdigit() else None
        for p in paras:
            style = p.get("style", "")
            text = p.get("text", "").strip()
            if not text:
                continue
            if "Heading" in style and section_pattern.match(text):
                in_section = True
                continue
            elif "Heading" in style and next_section and text.startswith(f"{next_section}."):
                break
            elif "Heading" in style and in_section and any(text.startswith(f"{i}.") for i in range(1, 15)):
                break
            if in_section and text and not _STRIP_RE.search(text):
                raw_lines.append(text)
        return raw_lines[:20]

    # Legend page — scale to fit within page margins
    doc.add_heading("Diagram Legend and Conventions", level=1)
    legend_path = PNG_DIR / "00-legend-conventions.png"
    if legend_path.exists():
        add_figure(doc, legend_path, "Diagram Legend and Conventions", width=Inches(6.0))
    doc.add_page_break()

    # Generate sections
    for sec in SRA_SECTIONS:
        doc.add_heading(f"§{sec['num']} — {sec['title']}", level=1)

        # Compliance alignment table with cross-reference rows
        cross_refs = []
        dep_nums = SRA_SECTION_DEPS.get(sec["num"], [])
        if dep_nums:
            dep_labels = []
            for dn in dep_nums:
                for s in SRA_SECTIONS:
                    if s["num"] == dn:
                        dep_labels.append(f"§{dn} {s['title']}")
                        break
            cross_refs.append(("Related SRA Sections", ", ".join(dep_labels)))
        pol_ids = SRA_TO_POLICIES.get(sec["num"], [])
        if pol_ids:
            cross_refs.append(("Associated Policies", ", ".join(pol_ids)))
        asr_refs = SRA_TO_ASR.get(sec["num"], [])
        if asr_refs:
            labels = [f"{dom} ({qr})" for dom, qr in asr_refs]
            cross_refs.append(("ASR Questionnaire Coverage", "; ".join(labels)))
        add_compliance_table(doc, sec, cross_refs=cross_refs)

        # Section prose
        content = extract_section_content(sec["num"], sra_paras) or extract_section_content(sec["num"], sra_anchored_paras)
        if content:
            for line in content:
                if len(line) <= 200:
                    doc.add_paragraph(line, style="List Bullet")
                else:
                    doc.add_paragraph(line)
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, "Section content derived from SRA V6 and governance memos. See source documents for full narrative.", italic=True)

        # Diagrams
        diag_keys = SEC_DIAGRAM_KEYS.get(sec["num"], [])
        for dk in diag_keys:
            png_file = DIAGRAM_MAP.get(dk, "")
            png_path = PNG_DIR / png_file
            if png_path.exists():
                # Strip leading 'NN-' number prefix from filename for caption
                stem = png_file.replace('.png', '').replace('-', ' ')
                fig_title = re.sub(r'^\d+\s+', '', stem).title()
                add_figure(doc, png_path, fig_title)

        doc.add_page_break()

    # Appendix A: Acronyms
    doc.add_heading("APPENDIX A — Acronyms and Definitions", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Term"
    tbl.rows[0].cells[1].text = "Definition"
    for term, defn in SRA_ACRONYMS:
        row = tbl.add_row()
        row.cells[0].text = term
        row.cells[1].text = defn
    style_table(tbl)

    doc.add_page_break()

    # Appendix B: Document Control
    doc.add_heading("APPENDIX B — Document Control", level=1)
    dc_path = PNG_DIR / "20-document-control-metadata.png"
    if dc_path.exists():
        add_figure(doc, dc_path, "Document Control Metadata")
    doc.add_paragraph("")
    tbl = doc.add_table(rows=7, cols=2)
    tbl.rows[0].cells[0].text = "Attribute"
    tbl.rows[0].cells[1].text = "Value"
    attrs = [
        ("Document Title", "Stride Security Reference Architecture"),
        ("Version", "8.0"),
        ("Date", "March 4, 2026"),
        ("Classification", "Confidential"),
        ("Author", "Security Architecture / GRC"),
        ("Supersedes", "SRA V7 (March 3, 2026)"),
    ]
    for i, (attr, val) in enumerate(attrs):
        tbl.rows[i + 1].cells[0].text = attr
        tbl.rows[i + 1].cells[1].text = val
    style_table(tbl)

    # Version history
    doc.add_paragraph("")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.rows[0].cells[0].text = "Version"
    tbl.rows[0].cells[1].text = "Date"
    tbl.rows[0].cells[2].text = "Change Description"
    tbl.rows[1].cells[0].text = "V6.0"
    tbl.rows[1].cells[1].text = "February 27, 2026"
    tbl.rows[1].cells[2].text = "Initial trust security architecture; 22 D2 diagrams; NIST CSF 2.0 alignment."
    tbl.rows[2].cells[0].text = "V7.0"
    tbl.rows[2].cells[1].text = "March 3, 2026"
    tbl.rows[2].cells[2].text = "Enhanced appendices (A–D); per-section FERPA/SOX alignment; diagram routing improvements."
    tbl.rows[3].cells[0].text = "V8.0"
    tbl.rows[3].cells[1].text = "March 4, 2026"
    tbl.rows[3].cells[2].text = "Cross-referencing between SRA sections, policies, and ASR questionnaire domains."
    style_table(tbl)

    doc.add_page_break()

    # Appendix C: NIST CSF 2.0 Mapping
    doc.add_heading("APPENDIX C — NIST CSF 2.0 Mapping and Implementation Status", level=1)
    doc.add_paragraph("Maps NIST CSF 2.0 subcategories to SRA sections and their implementation status.")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "CSF Subcategory"
    tbl.rows[0].cells[1].text = "Outcome Description"
    tbl.rows[0].cells[2].text = "SRA Section"
    tbl.rows[0].cells[3].text = "Status"
    for sub, desc, sra_sec, status in CSF_SRA_MAP:
        row = tbl.add_row()
        row.cells[0].text = sub
        row.cells[1].text = desc
        row.cells[2].text = sra_sec
        row.cells[3].text = status
        # Color the status cell
        if status == "Planned":
            set_cell_shading(row.cells[3], "FFEBEE")
        elif status == "In Progress":
            set_cell_shading(row.cells[3], "FFF3E0")
        else:
            set_cell_shading(row.cells[3], "E8F5E9")
    style_table(tbl)

    doc.add_page_break()

    # Appendix D: Evidence Traceability
    doc.add_heading("APPENDIX D — Evidence Traceability Matrix", level=1)
    doc.add_paragraph("Identifies primary evidence artifacts for each CSF subcategory, supporting audit readiness.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.rows[0].cells[0].text = "CSF Subcategory"
    tbl.rows[0].cells[1].text = "Evidence Artifacts"
    tbl.rows[0].cells[2].text = "Evidence Owner"
    for sub, evidence, owner in EVIDENCE_MAP:
        row = tbl.add_row()
        row.cells[0].text = sub
        row.cells[1].text = evidence
        row.cells[2].text = owner
    style_table(tbl)

    doc.add_page_break()

    # Appendix E: Compliance Gap and ASR Integration
    doc.add_heading("APPENDIX E — Compliance Gap Analysis (ASR Integration)", level=1)
    doc.add_paragraph(
        "This appendix identifies NIST CSF 2.0 subcategories, FERPA requirements, and SOX §404 "
        "ITGC controls where the Security Reference Architecture has partial or planned coverage. "
        "Gaps identified through the ASR questionnaire process are also cross-referenced."
    )

    doc.add_heading("NIST CSF 2.0 Subcategories — Planned or In Progress", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "CSF Subcategory"
    tbl.rows[0].cells[1].text = "Status"
    tbl.rows[0].cells[2].text = "SRA Section"
    tbl.rows[0].cells[3].text = "ASR Domain Relevance"
    csf_asr_domain_map = {
        "GV.OC": "Governance", "GV.PO": "Governance", "GV.OV": "Governance", "GV.RM": "Governance",
        "GV.RR": "Governance", "GV.SC": "Third-Party",
        "PR.AA": "IAM", "PR.DS": "Data Protection", "PR.PS": "Secure Dev", "PR.AT": "Governance",
        "PR.IR": "IR/BC",
        "ID.AM": "Governance", "ID.RA": "Vuln Mgmt", "ID.IM": "Vuln Mgmt",
        "DE.CM": "Vuln Mgmt", "DE.AE": "Vuln Mgmt",
        "RS.MA": "IR/BC", "RS.AN": "IR/BC", "RS.CO": "IR/BC",
        "RC.RP": "IR/BC",
    }
    for sub, desc, sra_sec, status in CSF_SRA_MAP:
        if status in ("Planned", "In Progress"):
            row = tbl.add_row()
            row.cells[0].text = sub
            row.cells[1].text = status
            row.cells[2].text = sra_sec
            cat_prefix = sub.split("-")[0] if "-" in sub else sub.split(".")[0]
            row.cells[3].text = csf_asr_domain_map.get(cat_prefix, "—")
            if status == "Planned":
                set_cell_shading(row.cells[1], "FFEBEE")
            else:
                set_cell_shading(row.cells[1], "FFF3E0")
    style_table(tbl)

    doc.add_heading("ASR-Identified Policy Gaps Impacting SRA", level=2)
    doc.add_paragraph(
        "The following ASR questionnaire topics address security controls not yet explicitly "
        "documented in the Stride policy framework. These gaps affect the SRA's ability to "
        "demonstrate full coverage across all control domains."
    )
    sra_gap_items = [
        ("Enterprise Architecture Governance", "GV.OV, GV.PO", "§2 Trust Architecture",
         "No formal EA review gate for application onboarding. ASR Question: architecture review process."),
        ("Business Necessity Classification", "ID.AM", "§9 Operational Trust",
         "No taxonomy or standard defining KTLO / business need / nice-to-have. ASR Question: business necessity."),
        ("Legal / IP Review Gate", "GV.RM, GV.PO", "§1 Executive Summary",
         "No published requirement for legal counsel review of application terms and IP. ASR Question: legal review."),
        ("Phased Deployment Framework", "GV.SC, PR.PS", "§11 Roadmap",
         "Crawl/walk/run model exists for AI (IISP 5.0) but not generalized. ASR Question: phased deployment."),
        ("Out-of-Band Patch Management", "PR.PS", "§8 Application Trust",
         "ISP 4.4 covers enterprise patching; applications with CLI/vendor-pushed updates lack specific controls."),
        ("Application Usage Telemetry / KPIs", "DE.CM, GV.OV", "§12 KPIs",
         "ISP 5.1 covers security monitoring but not application usage KPIs. ASR Question: usage monitoring."),
    ]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "Gap Topic"
    tbl.rows[0].cells[1].text = "CSF Alignment"
    tbl.rows[0].cells[2].text = "SRA Section"
    tbl.rows[0].cells[3].text = "Details"
    for topic, csf, sra_sec, detail in sra_gap_items:
        row = tbl.add_row()
        row.cells[0].text = topic
        row.cells[1].text = csf
        row.cells[2].text = sra_sec
        row.cells[3].text = detail
    style_table(tbl)

    ensure_copyable(doc)
    out_path = OUTPUT / "Stride_SRA_V8.docx"
    doc.save(str(out_path))
    print(f"  SUCCESS: {out_path} ({out_path.stat().st_size:,} bytes)")
    return out_path


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 3: ASR Questionnaire V4
# ══════════════════════════════════════════════════════════════════════

def build_asr_questionnaire_docx():
    print("Building DOCX: Stride ASR Questionnaire V4...")
    doc = new_doc()
    make_landscape(doc)

    add_title_page(doc, "Stride Application Security Review\nQuestionnaire", "Version 4.0",
                   extra_lines=[
                       "Discrete-choice questionnaire for application security risk assessment.",
                       "Organized by policy domain with regulatory cross-references.",
                       "Frameworks: NIST CSF 2.0 | FERPA (34 CFR Part 99) | SOX §404 ITGC",
                       "",
                       "Supersedes: ASR Questionnaire V3",
                   ])

    add_header_footer(doc, "Stride ASR Questionnaire V4")

    # Instructions
    doc.add_heading("Instructions", level=1)
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(
        "This questionnaire is used to assess the security posture of applications within the "
        "Stride technology portfolio. Each question maps to specific Stride ISP/IISP policies, "
        "NIST CSF 2.0 subcategories, and where applicable, FERPA and SOX §404 ITGC requirements."
    )

    doc.add_heading("How to Complete", level=2)
    instructions = [
        "For each question, select exactly one response that best describes the current state.",
        "Mark your selection with ✓ or X in the checkbox column.",
        "Questions marked Critical require immediate attention if the response indicates a gap.",
        "If a question is not applicable, select \"N/A\" and provide a brief justification.",
    ]
    for inst in instructions:
        doc.add_paragraph(inst, style="List Number")

    doc.add_heading("Scoring Model", level=2)

    # Weight tier table
    p = doc.add_paragraph()
    add_formatted_run(p, "Weight Tiers:", bold=True)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.rows[0].cells[0].text = "Weight Tier"
    tbl.rows[0].cells[1].text = "Priority"
    tbl.rows[0].cells[2].text = "Description"
    tbl.rows[0].cells[3].text = "Escalation"

    tiers = [
        ("Critical", "Highest", "Fundamental control — gap requires immediate remediation", "CISO notification", CRIT_BG_HEX),
        ("High", "High", "Important control — gap requires 30-day remediation", "Security Architecture review", HIGH_BG_HEX),
        ("Medium", "Moderate", "Supporting control — address in next planning cycle", "Risk acceptance option", MED_BG_HEX),
        ("Info", "Low", "Contextual / informational only", "None", None),
    ]
    for tier_name, priority, desc, escalation, bg_hex in tiers:
        row = tbl.add_row()
        row.cells[0].text = tier_name
        row.cells[1].text = priority
        row.cells[2].text = desc
        row.cells[3].text = escalation
        if bg_hex:
            set_cell_shading(row.cells[0], bg_hex)
    style_table(tbl)

    p = doc.add_paragraph()
    add_formatted_run(p, "Composite Scoring: ", bold=True)
    p.add_run("Weighted measurements are combined using a proprietary composite function "
              "that emphasizes the most severe finding and progressively discounts lesser ones.")

    p = doc.add_paragraph()
    add_formatted_run(p, "Interpretation: ", bold=True)
    p.add_run("0–25% = strong posture; 26–50% = adequate; 51–75% = elevated risk; "
              "76–100% = critical gaps requiring immediate action.")

    # Quick Links to Domain Sections
    doc.add_page_break()
    doc.add_heading("Quick Links — Assessment Domains", level=1)
    doc.add_paragraph(
        "Select a domain to navigate directly to its assessment questions."
    )
    for dom_entry in ASR_DOMAINS:
        p = doc.add_paragraph(style="List Bullet")
        add_formatted_run(p, dom_entry["domain"], bold=True, color=STRIDE_BLUE)

    # Application Info table
    doc.add_page_break()
    doc.add_heading("Application Information", level=1)
    tbl = doc.add_table(rows=8, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    fields = ["Application Name", "Business Owner", "Technical Owner",
              "CMDB ID", "Risk Classification", "Assessment Date", "Assessor"]
    for i, field in enumerate(fields):
        tbl.rows[i + 1].cells[0].text = field
        tbl.rows[i + 1].cells[1].text = ""
    style_table(tbl)

    doc.add_page_break()

    # Domain sections with question tables
    q_num = 0
    for domain in ASR_DOMAINS:
        doc.add_heading(domain["domain"], level=1)

        # Policy / CSF / Regulatory reference table
        ref_tbl = doc.add_table(rows=1, cols=2)
        ref_tbl.rows[0].cells[0].text = "Reference Type"
        ref_tbl.rows[0].cells[1].text = "Applicable References"

        row = ref_tbl.add_row()
        row.cells[0].text = "NIST CSF 2.0 Subcategories"
        row.cells[1].text = ", ".join(domain["csf_refs"])

        # Resolve policy titles for context
        ref_titles = []
        for pref in domain["policy_refs"]:
            for pol in POLICY_INVENTORY:
                if pol["id"] == pref:
                    ref_titles.append(f"{pref}: {pol['title']}")
                    break
        if ref_titles:
            row = ref_tbl.add_row()
            row.cells[0].text = "Policy Scope"
            # Each policy on its own line within the cell
            cell = row.cells[1]
            cell.text = ""  # clear default
            for t_idx, rt in enumerate(ref_titles):
                if t_idx == 0:
                    cell.paragraphs[0].text = rt
                else:
                    cell.add_paragraph(rt)

        # Regulatory notes as rows
        if "ferpa_note" in domain:
            row = ref_tbl.add_row()
            row.cells[0].text = "FERPA (34 CFR Part 99)"
            row.cells[1].text = domain["ferpa_note"]
        if "sox_note" in domain:
            row = ref_tbl.add_row()
            row.cells[0].text = "SOX §404 ITGC"
            row.cells[1].text = domain["sox_note"]

        style_table(ref_tbl)

        # Questions table: columns = #, Weight, Question, Choices, ✓, Notes
        n_questions = len(domain["questions"])
        tbl = doc.add_table(rows=1, cols=6)
        tbl.rows[0].cells[0].text = "#"
        tbl.rows[0].cells[1].text = "Weight"
        tbl.rows[0].cells[2].text = "Question"
        tbl.rows[0].cells[3].text = "Answer Choices"
        tbl.rows[0].cells[4].text = "✓"
        tbl.rows[0].cells[5].text = "Notes"

        for question in domain["questions"]:
            q_num += 1
            weight = question["weight"]
            n_c = len(question["choices"])
            scores = choice_scores(n_c)

            # One row per choice, merge # / Weight / Question cells vertically
            first_row_idx = len(tbl.rows)
            for i, choice in enumerate(question["choices"]):
                row = tbl.add_row()
                if i == 0:
                    row.cells[0].text = f"Q{q_num}"
                    row.cells[1].text = weight
                    row.cells[2].text = question["q"]
                # Choice text with score in smaller gray font
                cell3 = row.cells[3]
                cell3.text = ""  # clear default
                p3 = cell3.paragraphs[0]
                p3.add_run(choice)
                score_run = p3.add_run(f"  [{scores[i]}]")
                score_run.font.size = Pt(7)
                score_run.font.color.rgb = STRIDE_GRAY
                row.cells[4].text = "☐"
                if i == 0:
                    row.cells[5].text = ""

                # Color the weight cell
                if weight == "Critical":
                    set_cell_shading(row.cells[1], CRIT_BG_HEX)
                elif weight == "High":
                    set_cell_shading(row.cells[1], HIGH_BG_HEX)
                elif weight == "Medium":
                    set_cell_shading(row.cells[1], MED_BG_HEX)

            # Merge cells vertically for #, Weight, Question, Notes
            if n_c > 1:
                last_row_idx = len(tbl.rows) - 1
                for col_idx in [0, 1, 2, 5]:
                    tbl.cell(first_row_idx, col_idx).merge(tbl.cell(last_row_idx, col_idx))

        style_table(tbl)
        doc.add_page_break()

    # Assessment Summary
    doc.add_heading("Assessment Summary", level=1)
    doc.add_heading("Domain Summary", level=2)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.rows[0].cells[0].text = "Domain"
    tbl.rows[0].cells[1].text = "# Questions"
    tbl.rows[0].cells[2].text = "Score %"
    tbl.rows[0].cells[3].text = "Rating"
    tbl.rows[0].cells[4].text = "Notes"

    total_q = 0
    for domain in ASR_DOMAINS:
        row = tbl.add_row()
        row.cells[0].text = domain["domain"]
        row.cells[1].text = str(len(domain["questions"]))
        total_q += len(domain["questions"])

    row = tbl.add_row()
    row.cells[0].text = "OVERALL"
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    row.cells[1].text = str(total_q)
    style_table(tbl)

    # Rating scale
    doc.add_heading("Risk Rating Scale", level=2)
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "Norm %"
    tbl.rows[0].cells[1].text = "Description"
    tbl.rows[0].cells[2].text = "Rating"
    ratings = [
        ("0–25%", "Strong posture — controls are mature and effective", "Low"),
        ("26–50%", "Adequate posture — minor gaps with compensating controls", "Moderate"),
        ("51–75%", "Elevated risk — material gaps requiring remediation plans", "Elevated"),
        ("76–100%", "Critical risk — fundamental controls missing or ineffective", "Critical"),
    ]
    colors = [("E8F5E9", None), (None, None), ("FFF3E0", None), ("FFEBEE", None)]
    for i, ((pct, desc, rating), (bg, _)) in enumerate(zip(ratings, colors)):
        row = tbl.add_row()
        row.cells[0].text = pct
        row.cells[1].text = desc
        row.cells[2].text = rating
        if bg:
            set_row_shading(row, bg)
    style_table(tbl)

    # Assessor Certification
    doc.add_heading("Assessor Certification", level=2)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.rows[0].cells[0].text = "Assessor Signature:"
    tbl.rows[0].cells[1].text = "________________________________"
    tbl.rows[1].cells[0].text = "Date:"
    tbl.rows[1].cells[1].text = "________________________________"
    tbl.rows[2].cells[0].text = "Reviewer Signature:"
    tbl.rows[2].cells[1].text = "________________________________"
    tbl.rows[3].cells[0].text = "Date:"
    tbl.rows[3].cells[1].text = "________________________________"
    style_table(tbl, header_hex="FFFFFF")

    ensure_copyable(doc)
    out_path = OUTPUT / "Stride_ASR_Questionnaire_V4.docx"
    doc.save(str(out_path))
    print(f"  SUCCESS: {out_path} ({out_path.stat().st_size:,} bytes)")
    return out_path


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT 4: Governance Memos
# ══════════════════════════════════════════════════════════════════════

def _parse_markdown_to_docx(doc, md_text):
    """Parse markdown text and add it to a Word document with proper formatting.

    Handles: headings (##), bold (**), italic (*), tables (|), bullet lists (-),
    numbered lists (1.), blockquotes (>), and horizontal rules (---).
    """
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            # Just add a light separator paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, "—" * 40, color=STRIDE_GRAY, size=Pt(8))
            i += 1
            continue

        # Heading
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            level = min(level, 4)
            text = stripped.lstrip('#').strip()
            # Remove markdown bold markers
            text = text.replace('**', '')
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            if len(table_lines) >= 2:
                # Filter out separator rows (|---|---|)
                data_rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]  # skip empty first/last
                    if cells and not all(re.match(r'^[-:]+$', c) for c in cells):
                        data_rows.append(cells)

                if data_rows:
                    n_cols = max(len(r) for r in data_rows)
                    tbl = doc.add_table(rows=len(data_rows), cols=n_cols)
                    for row_idx, row_cells in enumerate(data_rows):
                        for col_idx, cell_text in enumerate(row_cells):
                            if col_idx < n_cols:
                                # Strip markdown bold
                                cell_text = cell_text.replace('**', '')
                                tbl.rows[row_idx].cells[col_idx].text = cell_text
                    style_table(tbl)
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            text = text.replace('**', '')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_run(p, text, italic=True, color=STRIDE_GRAY)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = _apply_inline_formatting(doc, text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            _apply_inline_formatting(doc, text, style="List Number")
            i += 1
            continue

        # Bold metadata line (key-value like **Field:** Value)
        bold_match = re.match(r'^\*\*(.+?)\*\*\s*(.*)', stripped)
        if bold_match:
            p = doc.add_paragraph()
            add_formatted_run(p, bold_match.group(1), bold=True)
            if bold_match.group(2):
                p.add_run(f" {bold_match.group(2)}")
            i += 1
            continue

        # Regular paragraph
        _apply_inline_formatting(doc, stripped)
        i += 1


def _apply_inline_formatting(doc, text, style=None):
    """Add a paragraph with inline bold/italic markdown converted to runs."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()

    # Split on **bold** and *italic* markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_formatted_run(p, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            add_formatted_run(p, part[1:-1], italic=True)
        else:
            p.add_run(part)
    return p


def build_governance_memos_docx():
    """Build .docx versions of all governance memos from .md sources."""
    print("Building DOCX: Governance Memos...")
    memo_files = sorted(glob.glob(str(MEMO_DIR / "GOV_MEMO_SRA_S*.md")))

    if not memo_files:
        print("  No governance memo .md files found.")
        return []

    results = []
    for md_path in memo_files:
        md_path = Path(md_path)
        stem = md_path.stem  # e.g. GOV_MEMO_SRA_S1_Executive_Summary
        print(f"  Processing {stem}...")

        md_text = md_path.read_text(errors='replace')

        doc = new_doc()
        make_landscape(doc)

        # Extract title from first heading
        title_match = re.search(r'^#\s+(.+)', md_text, re.MULTILINE)
        title = title_match.group(1) if title_match else stem.replace('_', ' ')
        title = title.replace('**', '')

        add_header_footer(doc, f"Stride — {title}")

        # Parse the markdown content
        _parse_markdown_to_docx(doc, md_text)

        ensure_copyable(doc)
        out_name = f"{stem}.docx"
        out_path = OUTPUT / out_name
        doc.save(str(out_path))
        size = out_path.stat().st_size
        print(f"    → {out_path.name} ({size:,} bytes)")
        results.append(out_path)

    print(f"  SUCCESS: {len(results)} governance memos built")
    return results


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("STRIDE DELIVERABLES — WORD (.docx) BUILD")
    print(f"Date: {date.today()}")
    print("=" * 60)

    d1 = build_policy_cross_reference_docx()
    d2 = build_sra_v8_docx()
    d3 = build_asr_questionnaire_docx()
    memos = build_governance_memos_docx()

    print("\n" + "=" * 60)
    print("BUILD COMPLETE")
    print("=" * 60)

    all_docs = [
        ("Policy Cross Reference V7", d1),
        ("SRA V8", d2),
        ("ASR Questionnaire V4", d3),
    ]
    for name, path in all_docs:
        if path and path.exists():
            print(f"  ✓ {name}: {path.name} ({path.stat().st_size:,} bytes)")
        else:
            print(f"  ✗ {name}: FAILED")

    if memos:
        print(f"  ✓ Governance Memos: {len(memos)} documents")
        for m in memos:
            print(f"      {m.name} ({m.stat().st_size:,} bytes)")

    print(f"\nOutput directory: {OUTPUT}")
