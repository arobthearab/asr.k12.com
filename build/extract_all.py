#!/usr/bin/env python3
"""Extract text from all Stride policy PDFs and Work Papers docx files.
Outputs JSON to build/extracted/
"""
import json, os, sys, pathlib

import pypdf
import docx as python_docx

BASE = pathlib.Path(__file__).resolve().parent          # <repo>/build/
REPO = BASE.parent                                      # <repo>/
OUT = BASE / "extracted"
OUT.mkdir(exist_ok=True)

# --- Extract PDFs ---
pdf_dir = REPO / "policies"
pdf_results = {}
for pdf_path in sorted(pdf_dir.glob("*.pdf")):
    name = pdf_path.stem
    try:
        reader = pypdf.PdfReader(str(pdf_path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)
        full_text = "\n\n--- PAGE BREAK ---\n\n".join(pages)
        pdf_results[name] = {
            "file": pdf_path.name,
            "pages": len(reader.pages),
            "chars": len(full_text),
            "text": full_text
        }
        print(f"  OK: {name} ({len(reader.pages)} pp, {len(full_text)} chars)")
    except Exception as e:
        pdf_results[name] = {"file": pdf_path.name, "error": str(e)}
        print(f"  FAIL: {name}: {e}")

with open(OUT / "policies.json", "w") as f:
    json.dump(pdf_results, f, indent=2)
print(f"\nPolicies: {len(pdf_results)} files -> {OUT / 'policies.json'}")

# --- Extract DOCX files ---
docx_dir = REPO / "workpapers"
docx_results = {}
for docx_path in sorted(docx_dir.glob("*.docx")):
    name = docx_path.stem
    try:
        doc = python_docx.Document(str(docx_path))
        paragraphs = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            paragraphs.append({
                "style": style_name,
                "text": p.text
            })
        # Also extract tables
        tables = []
        for ti, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        
        full_text = "\n".join(p["text"] for p in paragraphs)
        docx_results[name] = {
            "file": docx_path.name,
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "chars": len(full_text),
            "structured": paragraphs,
            "table_data": tables
        }
        print(f"  OK: {name} ({len(paragraphs)} paras, {len(tables)} tables, {len(full_text)} chars)")
    except Exception as e:
        docx_results[name] = {"file": docx_path.name, "error": str(e)}
        print(f"  FAIL: {name}: {e}")

with open(OUT / "workpapers.json", "w") as f:
    json.dump(docx_results, f, indent=2)
print(f"\nWork Papers: {len(docx_results)} files -> {OUT / 'workpapers.json'}")

# --- Extract MD files ---
md_results = {}
for md_path in sorted(docx_dir.glob("*.md")):
    name = md_path.stem
    text = md_path.read_text(encoding="utf-8", errors="replace")
    md_results[name] = {"file": md_path.name, "chars": len(text), "text": text}
    print(f"  OK: {name} ({len(text)} chars)")

with open(OUT / "markdown.json", "w") as f:
    json.dump(md_results, f, indent=2)
print(f"\nMarkdown: {len(md_results)} files -> {OUT / 'markdown.json'}")

# --- Summary ---
print("\n=== EXTRACTION SUMMARY ===")
ok_pdfs = sum(1 for v in pdf_results.values() if "text" in v)
fail_pdfs = sum(1 for v in pdf_results.values() if "error" in v)
ok_docx = sum(1 for v in docx_results.values() if "structured" in v)
fail_docx = sum(1 for v in docx_results.values() if "error" in v)
print(f"PDFs:  {ok_pdfs} OK, {fail_pdfs} failed")
print(f"DOCX:  {ok_docx} OK, {fail_docx} failed")
print(f"MD:    {len(md_results)} OK")
total_chars = sum(v.get("chars", 0) for v in pdf_results.values()) + \
              sum(v.get("chars", 0) for v in docx_results.values()) + \
              sum(v.get("chars", 0) for v in md_results.values())
print(f"Total chars extracted: {total_chars:,}")
